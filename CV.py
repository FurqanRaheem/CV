from docx import Document
import docx.enum.text as text
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm

import os
import pyinputplus as pyip
import datetime as dt


def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '35')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '#2E74BC')
    pBdr.append(bottom)


def addHeading(heading):
    head = doc.add_paragraph()
    run = head.add_run(heading)
    run.bold = True
    run.font.size = Pt(14)
    insertHR(head)


def getDates():
    start_month = pyip.inputMonth("Starting Month: ")
    start_year = str(pyip.inputDate("Starting Year: ", formats=["%Y"]))
    year, month, day = map(int, start_year.split('-'))
    date = dt.date(year, month, day)
    start_date = f"{start_month} {date.year}"

    end_month = pyip.inputMonth("Ending Month: ")
    end_year = str(pyip.inputDate("Ending Year: ", formats=["%Y"]))
    year, month, day = map(int, end_year.split('-'))
    date = dt.date(year, month, day)
    end_date = f"{end_month} {date.year}"

    return start_date, end_date


def setDates(startDate, endDate):
    date = doc.add_paragraph()
    date.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT
    run = date.add_run(f"{startDate} - {endDate}")
    dateFont = run.font
    dateFont.italic = True
    dateFont.color.rgb = RGBColor(0x2E, 0x74, 0xBC)


doc = Document()  # Create Document
margin = 1.27
doc.styles['Normal'].font.name = "Arial"
for section in doc.sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

# name = input("Name: ")
# number = input("Phone number: ")
# email = input("Email: ")
# address = input("Address: ")
name = "Furqan Khan"
number = "07746435680"
email = "furqanraheem10@gmail.com"
address = "191 Tennyson Place, Bradford BD3 0AE"

nameHeading = doc.add_paragraph()
nameHeading.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
run = nameHeading.add_run(f"{name}\n")
run.font.size = Pt(18)

contact_details = doc.add_paragraph()
contact_details.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
run = contact_details.add_run(f"{address}\n☏{number}✉{email}")
run.font.size = Pt(12)


# Personal Statement
addHeading("Personal Statement")
# doc.add_paragraph(input("Personal Statement:\n"))
doc.add_paragraph(
    "An A-levels student with a proven leadership ability in an educational, fitness and professional setting. Being resilient and confident, I work well under pressure and remain positive in any situation I face, thriving off challenges. I provide timely, efficient and accurate support where required to meet the business needs. I have a friendly personality, an inquiring mind and a hunger for acquiring knowledge. I always aim to give 100% to any given task and remain motivated to work well within a team and on my own initiative. I am currently seeking to move up in my career, take on more responsibilities and learn new skills in a new environment. I am passionate to use my current skills and experience to help a business meet its goals and simultaneously broaden my own knowledge and skillset.")

# Key Skills
addHeading("Key Skills")
print("\n\nSkills. Type 'None' when finished.")
while True:
    skill = input("Skills: ")
    if skill.upper() != "NONE":
        if not skill == "":
            doc.add_paragraph(skill, style="List Bullet")  # Add Bullet Point
        else:
            print("\nEmpty bullet point")
            print("Type 'None' when finished\n")
    else:
        break

# Employment History
addHeading("Employment History")
print("Work experience")
while not pyip.inputYesNo("More work experience (Yes/No) ") == "no":
    company = input("Company: ")
    role = input("Role: ")
    city = input("City: ")
    startDate, endDate = getDates()

    font = doc.add_paragraph().add_run(f"{role}, {company}, {city}").font
    font.size = Pt(12)
    font.bold = True

    setDates(startDate, endDate)

    doc.add_paragraph("Responsibilities:")
    print("\nResponsibilities. Type 'None' when finished.")
    while True:
        responsibility = input("Responsibilities: ")
        if responsibility.upper() != "NONE":
            if not responsibility == "":
                doc.add_paragraph(responsibility, style="List Bullet")  # Add Bullet Point
        else:
            break

# Education
addHeading("Education")
print("Education")
while not pyip.inputYesNo("More education history (Yes/No) ") == "no":
    institution = input("Institution: ")
    course = input("Course: ")
    startDate, endDate = getDates()
    doc.add_paragraph().add_run(f"{institution}").bold = True
    setDates(startDate, endDate)

    if pyip.inputYesNo("Add Modules/Subjects (Yes/No)? ") == "yes":
        doc.add_paragraph(f"{course}:")
        print("\nSubjects.\nType 'None' when finished.")
        while True:
            subject = input("Subject: ")
            if subject.upper() != "NONE":
                if not subject == "":
                    doc.add_paragraph(subject, style="List Bullet")  # Add Bullet Point
            else:
                break

# Languages
if pyip.inputYesNo("Would you like to add a languages section? ") == "yes":
    addHeading("Languages")
    print("\nLanguages. Type 'None' when finished.")
    while True:
        subject = input("Language: ")
        if subject.upper() != "NONE":
            if not subject == "":
                doc.add_paragraph(subject, style="List Bullet")  # Add Bullet Point
        else:
            break

# Save Document

doc.save(f"{os.getcwd()}\\cv.docx")
